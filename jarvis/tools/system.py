"""System tools — time, weather, open apps, reminders, file operations, music, summarizer."""

import os
import re
import subprocess
import datetime
import json
import urllib.parse
from pathlib import Path

import httpx

from jarvis.config import DEFAULT_CITY, ANTHROPIC_API_KEY
from jarvis.tools.base import Tool, registry

REMINDERS_FILE = Path(__file__).parent.parent.parent / "data" / "reminders.json"


def get_current_time(**kwargs) -> str:
    now = datetime.datetime.now()
    return now.strftime("%A, %B %d, %Y at %I:%M %p")


WEATHER_CODES = {
    0: "clear sky", 1: "mainly clear", 2: "partly cloudy", 3: "overcast",
    45: "foggy", 48: "depositing rime fog",
    51: "light drizzle", 53: "moderate drizzle", 55: "dense drizzle",
    61: "slight rain", 63: "moderate rain", 65: "heavy rain",
    71: "slight snow", 73: "moderate snow", 75: "heavy snow",
    80: "slight rain showers", 81: "moderate rain showers", 82: "violent rain showers",
    95: "thunderstorm", 96: "thunderstorm with slight hail", 99: "thunderstorm with heavy hail",
}


def get_weather(city: str = "", when: str = "today") -> str:
    """Fetch weather from Open-Meteo. Supports today, tomorrow, or multi-day forecast."""
    if not city:
        city = DEFAULT_CITY
    # Try full name first, then just the city part (API doesn't like "City, State")
    geo = httpx.get(
        "https://geocoding-api.open-meteo.com/v1/search",
        params={"name": city, "count": 1},
    ).json()

    if not geo.get("results") and "," in city:
        city_only = city.split(",")[0].strip()
        geo = httpx.get(
            "https://geocoding-api.open-meteo.com/v1/search",
            params={"name": city_only, "count": 1},
        ).json()

    if not geo.get("results"):
        return f"Could not find weather data for {city}."

    loc = geo["results"][0]
    lat, lon = loc["latitude"], loc["longitude"]

    weather = httpx.get(
        "https://api.open-meteo.com/v1/forecast",
        params={
            "latitude": lat,
            "longitude": lon,
            "current": "temperature_2m,weathercode,windspeed_10m,relative_humidity_2m",
            "daily": "weathercode,temperature_2m_max,temperature_2m_min,precipitation_probability_max,windspeed_10m_max",
            "temperature_unit": "fahrenheit",
            "windspeed_unit": "mph",
            "forecast_days": 7,
        },
    ).json()

    lines = []

    if when == "today" or when == "now":
        cur = weather.get("current", {})
        temp = cur.get("temperature_2m", "?")
        wind = cur.get("windspeed_10m", "?")
        humidity = cur.get("relative_humidity_2m", "?")
        desc = WEATHER_CODES.get(cur.get("weathercode", 0), "unknown")
        lines.append(f"Right now in {loc['name']}: {temp}F, {desc}, wind {wind} mph, humidity {humidity}%.")

        # Also include today's high/low
        daily = weather.get("daily", {})
        if daily.get("temperature_2m_max"):
            hi = daily["temperature_2m_max"][0]
            lo = daily["temperature_2m_min"][0]
            rain = daily.get("precipitation_probability_max", [0])[0]
            lines.append(f"Today's forecast: high {hi}F, low {lo}F, {rain}% chance of rain.")

    elif when == "tomorrow":
        daily = weather.get("daily", {})
        if daily.get("temperature_2m_max") and len(daily["temperature_2m_max"]) > 1:
            hi = daily["temperature_2m_max"][1]
            lo = daily["temperature_2m_min"][1]
            code = daily.get("weathercode", [0, 0])[1]
            desc = WEATHER_CODES.get(code, "unknown")
            rain = daily.get("precipitation_probability_max", [0, 0])[1]
            wind = daily.get("windspeed_10m_max", [0, 0])[1]
            date = daily["time"][1]
            lines.append(f"Tomorrow ({date}) in {loc['name']}: {desc}.")
            lines.append(f"High {hi}F, low {lo}F, wind up to {wind} mph, {rain}% chance of rain.")

    elif when == "week":
        daily = weather.get("daily", {})
        lines.append(f"7-day forecast for {loc['name']}:")
        for i in range(min(7, len(daily.get("time", [])))):
            date = daily["time"][i]
            hi = daily["temperature_2m_max"][i]
            lo = daily["temperature_2m_min"][i]
            code = daily.get("weathercode", [0])[i]
            desc = WEATHER_CODES.get(code, "unknown")
            rain = daily.get("precipitation_probability_max", [0])[i]
            lines.append(f"  {date}: {desc}, {hi}F/{lo}F, {rain}% rain")

    else:
        # Default to today
        return get_weather(city, "today")

    return "\n".join(lines)


def open_application(name: str) -> str:
    """Open an application by name on Windows."""
    app_map = {
        "chrome": "chrome",
        "google chrome": "chrome",
        "browser": "chrome",
        "notepad": "notepad",
        "calculator": "calc",
        "file explorer": "explorer",
        "explorer": "explorer",
        "terminal": "wt",
        "command prompt": "cmd",
        "cmd": "cmd",
        "spotify": "spotify",
        "discord": "discord",
        "slack": "slack",
        "vscode": "code",
        "vs code": "code",
        "visual studio code": "code",
    }

    key = name.lower().strip()
    exe = app_map.get(key, key)

    try:
        subprocess.Popen(
            f"start {exe}", shell=True,
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
        )
        return f"Opening {name}."
    except Exception as e:
        return f"Failed to open {name}: {e}"


def _send_ntfy(title: str, message: str):
    """Send push notification via ntfy.sh."""
    try:
        httpx.post(
            "https://ntfy.sh/kalshi-trader-alerts",
            content=message.encode("utf-8"),
            headers={"Title": title, "Priority": "high", "Tags": "bell"},
            timeout=10,
        )
    except Exception:
        pass


def send_notification(message: str = "", title: str = "JARVIS", **kwargs) -> str:
    """Send a push notification to Deagz's phone via ntfy."""
    try:
        resp = httpx.post(
            "https://ntfy.sh/kalshi-trader-alerts",
            content=message.encode("utf-8"),
            headers={"Title": title, "Priority": "default", "Tags": "robot"},
            timeout=10,
        )
        if resp.status_code == 200:
            return f"Notification sent to your phone: {message}"
        return f"Failed to send notification: {resp.status_code}"
    except Exception as e:
        return f"Notification failed: {e}"


def set_reminder(message: str, minutes: int = 0) -> str:
    """Save a reminder. If minutes > 0, sends a push notification after the delay."""
    import threading

    REMINDERS_FILE.parent.mkdir(parents=True, exist_ok=True)

    reminders = []
    if REMINDERS_FILE.exists():
        reminders = json.loads(REMINDERS_FILE.read_text())

    now = datetime.datetime.now()
    reminder = {
        "message": message,
        "created": now.isoformat(),
        "due": (now + datetime.timedelta(minutes=minutes)).isoformat() if minutes > 0 else None,
    }
    reminders.append(reminder)
    REMINDERS_FILE.write_text(json.dumps(reminders, indent=2))

    if minutes > 0:
        # Schedule ntfy push notification
        def _notify():
            import time
            time.sleep(minutes * 60)
            _send_ntfy("JARVIS Reminder", message)

        threading.Thread(target=_notify, daemon=True).start()
        return f"Reminder set: '{message}' in {minutes} minutes. I'll ping your phone."
    else:
        # Immediate reminder — send now
        _send_ntfy("JARVIS Reminder", message)
        return f"Reminder saved and sent to your phone: '{message}'."


def list_reminders() -> str:
    """List all saved reminders."""
    if not REMINDERS_FILE.exists():
        return "No reminders set."

    reminders = json.loads(REMINDERS_FILE.read_text())
    if not reminders:
        return "No reminders set."

    lines = []
    for i, r in enumerate(reminders, 1):
        due = f" (due: {r['due']})" if r.get("due") else ""
        lines.append(f"{i}. {r['message']}{due}")
    return "\n".join(lines)


def web_search(query: str) -> str:
    """Search the web using DuckDuckGo instant answers (free, no key)."""
    resp = httpx.get(
        "https://api.duckduckgo.com/",
        params={"q": query, "format": "json", "no_html": 1},
        timeout=10,
    ).json()

    abstract = resp.get("AbstractText", "")
    if abstract:
        return abstract

    # Try related topics
    related = resp.get("RelatedTopics", [])
    if related:
        results = []
        for topic in related[:3]:
            if isinstance(topic, dict) and "Text" in topic:
                results.append(topic["Text"])
        if results:
            return " | ".join(results)

    return f"No instant answer found for '{query}'. Try asking me to search more specifically."


def run_command(command: str) -> str:
    """Run a shell command on the computer and return the output."""
    try:
        result = subprocess.run(
            command, shell=True, capture_output=True, text=True, timeout=30,
        )
        output = result.stdout.strip()
        if result.returncode != 0 and result.stderr:
            output += f"\nError: {result.stderr.strip()}"
        return output if output else "Command completed with no output."
    except subprocess.TimeoutExpired:
        return "Command timed out after 30 seconds."
    except Exception as e:
        return f"Failed to run command: {e}"


def open_url(url: str) -> str:
    """Open a URL in the default web browser."""
    try:
        if not url.startswith("http"):
            url = "https://" + url
        import webbrowser
        webbrowser.open(url)
        return f"Opening {url} in your browser."
    except Exception as e:
        return f"Failed to open URL: {e}"


def read_file(path: str) -> str:
    """Read the contents of a file."""
    try:
        p = Path(path).expanduser()
        if not p.exists():
            return f"File not found: {path}"
        content = p.read_text(encoding="utf-8", errors="replace")
        if len(content) > 2000:
            return content[:2000] + f"\n... (truncated, {len(content)} chars total)"
        return content
    except Exception as e:
        return f"Failed to read file: {e}"


def write_file(path: str, content: str) -> str:
    """Write content to a file. Creates the file if it doesn't exist."""
    try:
        p = Path(path).expanduser()
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(content, encoding="utf-8")
        return f"Written to {path}."
    except Exception as e:
        return f"Failed to write file: {e}"


def list_directory(path: str = ".") -> str:
    """List files and folders in a directory."""
    try:
        p = Path(path).expanduser()
        if not p.exists():
            return f"Directory not found: {path}"
        items = sorted(p.iterdir(), key=lambda x: (not x.is_dir(), x.name.lower()))
        lines = []
        for item in items[:50]:
            prefix = "[DIR] " if item.is_dir() else "      "
            size = ""
            if item.is_file():
                s = item.stat().st_size
                if s > 1_000_000:
                    size = f" ({s / 1_000_000:.1f} MB)"
                elif s > 1000:
                    size = f" ({s / 1000:.0f} KB)"
            lines.append(f"{prefix}{item.name}{size}")
        result = "\n".join(lines)
        if len(items) > 50:
            result += f"\n... and {len(items) - 50} more items"
        return result
    except Exception as e:
        return f"Failed to list directory: {e}"


def kill_process(name: str) -> str:
    """Kill a running process by name."""
    try:
        result = subprocess.run(
            f'taskkill /IM "{name}" /F',
            shell=True, capture_output=True, text=True,
        )
        if result.returncode == 0:
            return f"Killed {name}."
        return f"Could not kill {name}: {result.stderr.strip()}"
    except Exception as e:
        return f"Failed: {e}"


def get_system_info() -> str:
    """Get basic system information — CPU, memory, disk usage."""
    try:
        lines = []
        # CPU
        cpu = subprocess.run(
            'wmic cpu get name /value', shell=True, capture_output=True, text=True
        ).stdout.strip()
        for line in cpu.split("\n"):
            if "Name=" in line:
                lines.append(f"CPU: {line.split('=', 1)[1].strip()}")

        # Memory
        mem = subprocess.run(
            'wmic os get FreePhysicalMemory,TotalVisibleMemorySize /value',
            shell=True, capture_output=True, text=True,
        ).stdout.strip()
        total = free = 0
        for line in mem.split("\n"):
            if "TotalVisibleMemorySize=" in line:
                total = int(line.split("=")[1].strip()) / 1024 / 1024
            if "FreePhysicalMemory=" in line:
                free = int(line.split("=")[1].strip()) / 1024 / 1024
        if total > 0:
            used = total - free
            lines.append(f"RAM: {used:.1f} GB used / {total:.1f} GB total")

        # Disk
        disk = subprocess.run(
            'wmic logicaldisk where "DeviceID=\'C:\'" get FreeSpace,Size /value',
            shell=True, capture_output=True, text=True,
        ).stdout.strip()
        d_total = d_free = 0
        for line in disk.split("\n"):
            if "Size=" in line:
                d_total = int(line.split("=")[1].strip()) / 1024**3
            if "FreeSpace=" in line:
                d_free = int(line.split("=")[1].strip()) / 1024**3
        if d_total > 0:
            lines.append(f"Disk C: {d_free:.0f} GB free / {d_total:.0f} GB total")

        # Battery
        batt = subprocess.run(
            'wmic path win32_battery get EstimatedChargeRemaining /value',
            shell=True, capture_output=True, text=True,
        ).stdout.strip()
        for line in batt.split("\n"):
            if "EstimatedChargeRemaining=" in line:
                lines.append(f"Battery: {line.split('=')[1].strip()}%")

        return "\n".join(lines) if lines else "Could not retrieve system info."
    except Exception as e:
        return f"Failed: {e}"


def set_volume(level: int) -> str:
    """Set system volume (0-100)."""
    try:
        level = max(0, min(100, level))
        # Convert 0-100 to 0-65535
        val = int(level / 100 * 65535)
        subprocess.run(
            f'powershell -Command "(New-Object -ComObject WScript.Shell).SendKeys([char]173)"',
            shell=True, capture_output=True,
        )
        # Use nircmd if available, otherwise PowerShell
        subprocess.run(
            f'powershell -Command "Set-AudioDevice -PlaybackVolume {level}"',
            shell=True, capture_output=True,
        )
        return f"Volume set to {level}%."
    except Exception as e:
        return f"Failed to set volume: {e}"


def screenshot() -> str:
    """Take a screenshot and save it to the desktop."""
    try:
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        path = f"C:\\Users\\brian\\Desktop\\screenshot_{ts}.png"
        subprocess.run(
            f'powershell -Command "Add-Type -AssemblyName System.Windows.Forms; '
            f"[System.Windows.Forms.Screen]::PrimaryScreen | ForEach-Object {{ "
            f"$bmp = New-Object System.Drawing.Bitmap($_.Bounds.Width, $_.Bounds.Height); "
            f"$gfx = [System.Drawing.Graphics]::FromImage($bmp); "
            f"$gfx.CopyFromScreen($_.Bounds.Location, [System.Drawing.Point]::Empty, $_.Bounds.Size); "
            f"$bmp.Save('{path}'); "
            f'$gfx.Dispose(); $bmp.Dispose() }}"',
            shell=True, check=True, timeout=10,
        )
        return f"Screenshot saved to {path}."
    except Exception as e:
        return f"Failed to take screenshot: {e}"


def play_music(query: str, service: str = "spotify") -> str:
    """Play music on Spotify or YouTube."""
    import webbrowser
    try:
        encoded = urllib.parse.quote(query)
        if service.lower() == "youtube":
            url = f"https://www.youtube.com/results?search_query={encoded}"
            webbrowser.open(url)
            return f"Searching YouTube for '{query}'."
        else:
            url = f"https://open.spotify.com/search/{encoded}"
            webbrowser.open(url)
            return f"Searching Spotify for '{query}'."
    except Exception as e:
        return f"Failed to open music: {e}"


def control_music(action: str) -> str:
    """Control media playback — play/pause, next, previous, volume up/down."""
    key_map = {
        "play": "0xB3",       # Play/Pause
        "pause": "0xB3",      # Play/Pause (toggle)
        "next": "0xB0",       # Next Track
        "skip": "0xB0",       # Next Track
        "previous": "0xB1",   # Previous Track
        "back": "0xB1",       # Previous Track
        "volume_up": "0xAF",  # Volume Up
        "volume_down": "0xAE",  # Volume Down
        "mute": "0xAD",       # Mute
    }

    action_lower = action.lower().strip()
    vk_code = key_map.get(action_lower)

    if not vk_code:
        return f"Unknown action '{action}'. Try: play, pause, next, previous, volume_up, volume_down, mute."

    try:
        subprocess.run(
            f'powershell -Command "'
            f'$key = {vk_code}; '
            f'$hwnd = 0; '
            f'Add-Type -TypeDefinition @\\"\\nusing System; using System.Runtime.InteropServices;\\n'
            f'public class MediaKey {{ [DllImport(\\"user32.dll\\")] public static extern void keybd_event(byte bVk, byte bScan, uint dwFlags, UIntPtr dwExtraInfo); }}\\n\\"@;'
            f'[MediaKey]::keybd_event({vk_code}, 0, 0, [UIntPtr]::Zero); '
            f'[MediaKey]::keybd_event({vk_code}, 0, 2, [UIntPtr]::Zero)"',
            shell=True, capture_output=True, timeout=5,
        )
        return f"Media {action_lower} executed."
    except Exception as e:
        return f"Failed to control media: {e}"


def create_document(title: str = "Document", content: str = "", format: str = "docx", **kwargs) -> str:
    """Create a Word document or text file on the Desktop."""
    try:
        desktop = Path(os.path.expanduser("~/Desktop"))
        safe_title = re.sub(r'[<>:"/\\|?*]', '_', title)

        if format.lower() == "txt":
            path = desktop / f"{safe_title}.txt"
            path.write_text(content, encoding="utf-8")
            return f"Text file saved to {path}"

        # Word document
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add date
        doc.add_paragraph(
            datetime.datetime.now().strftime("%B %d, %Y"),
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")  # spacer

        # Content — split by newlines, detect headers with ##
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph("")
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)

        path = desktop / f"{safe_title}.docx"
        doc.save(str(path))

        # Open the file
        subprocess.Popen(["cmd", "/c", "start", "", str(path)], shell=False)
        return f"Document saved and opened: {path}"
    except Exception as e:
        return f"Failed to create document: {e}"


def get_news(topic: str = "", **kwargs) -> str:
    """Get latest news headlines."""
    try:
        query = topic if topic else "top news today"
        # Use DuckDuckGo news
        resp = httpx.get(
            "https://duckduckgo.com/",
            params={"q": f"{query} news", "format": "json", "t": "jarvis"},
            headers={"User-Agent": "Mozilla/5.0"},
            timeout=10,
        )

        # Fallback: use the web search + Claude to summarize headlines
        from duckduckgo_search import DDGS
        with DDGS() as ddgs:
            results = list(ddgs.news(query, max_results=5))

        if not results:
            return f"No news found for '{topic}'."

        lines = []
        for r in results:
            lines.append(f"- {r.get('title', 'N/A')} ({r.get('source', 'unknown')})")
        return f"Top headlines for '{topic or 'today'}':\n" + "\n".join(lines)
    except Exception:
        # Fallback to web search
        try:
            from duckduckgo_search import DDGS
            with DDGS() as ddgs:
                results = list(ddgs.text(f"{topic or 'top'} news today", max_results=5))
            if results:
                lines = [f"- {r.get('title', 'N/A')}" for r in results]
                return "Latest news:\n" + "\n".join(lines)
        except Exception as e2:
            return f"News unavailable: {e2}"


def draft_email(to: str = "", subject: str = "", body: str = "", **kwargs) -> str:
    """Compose an email and open it in the default email client."""
    import webbrowser
    try:
        params = {}
        if subject:
            params["subject"] = subject
        if body:
            params["body"] = body
        query = urllib.parse.urlencode(params, quote_via=urllib.parse.quote)
        mailto = f"mailto:{urllib.parse.quote(to)}"
        if query:
            mailto += f"?{query}"
        webbrowser.open(mailto)
        return f"Email draft opened — To: {to}, Subject: {subject}"
    except Exception as e:
        return f"Failed to open email: {e}"


def analyze_screenshot(**kwargs) -> str:
    """Take a screenshot and analyze it with Claude's vision."""
    import base64
    try:
        result = subprocess.run(
            'powershell -Command "Add-Type -AssemblyName System.Windows.Forms; '
            '[System.Windows.Forms.Screen]::PrimaryScreen.Bounds | '
            'ForEach-Object { $bmp = New-Object System.Drawing.Bitmap($_.Width, $_.Height); '
            '$g = [System.Drawing.Graphics]::FromImage($bmp); '
            '$g.CopyFromScreen($_.Location, [System.Drawing.Point]::Empty, $_.Size); '
            r'$path = \"$env:TEMP\jarvis_screen.png\"; '
            '$bmp.Save($path); $path }"',
            shell=True, capture_output=True, text=True, timeout=10,
        )
        img_path = result.stdout.strip()
        if not img_path or not Path(img_path).exists():
            return "Failed to capture screenshot."

        with open(img_path, "rb") as f:
            img_b64 = base64.b64encode(f.read()).decode()

        import anthropic
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        resp = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=300,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": "image/png",
                            "data": img_b64,
                        },
                    },
                    {
                        "type": "text",
                        "text": "Describe what's on this screen briefly. Focus on what the user is doing — the active window, any important content, or notable details. Keep it to 2-3 sentences.",
                    },
                ],
            }],
        )
        return resp.content[0].text
    except Exception as e:
        return f"Failed to analyze screen: {e}"


def set_alarm(minutes: int = 0, time_str: str = "", message: str = "Alarm") -> str:
    """Set an alarm — either in X minutes or at a specific time."""
    import threading
    import winsound

    def _ring(msg, delay):
        import time
        time.sleep(delay)
        # Beep 3 times
        for _ in range(3):
            winsound.Beep(1000, 500)
            time.sleep(0.3)
        # Also show a notification
        try:
            subprocess.run(
                f'powershell -Command "Add-Type -AssemblyName System.Windows.Forms; '
                f"[System.Windows.Forms.MessageBox]::Show('{msg}', 'JARVIS Alarm')\"",
                shell=True, timeout=10,
            )
        except Exception:
            pass

    if time_str:
        # Parse time like "7:00 AM", "7am", "19:00"
        now = datetime.datetime.now()
        try:
            for fmt in ("%I:%M %p", "%I:%M%p", "%I %p", "%I%p", "%H:%M"):
                try:
                    t = datetime.datetime.strptime(time_str.upper().strip(), fmt)
                    target = now.replace(hour=t.hour, minute=t.minute, second=0)
                    if target <= now:
                        target += datetime.timedelta(days=1)
                    delay = (target - now).total_seconds()
                    threading.Thread(target=_ring, args=(message, delay), daemon=True).start()
                    return f"Alarm set for {target.strftime('%I:%M %p')}: {message}"
                except ValueError:
                    continue
            return f"Could not parse time: {time_str}. Try format like '7:00 AM' or '19:00'."
        except Exception as e:
            return f"Failed to set alarm: {e}"
    elif minutes > 0:
        delay = minutes * 60
        threading.Thread(target=_ring, args=(message, delay), daemon=True).start()
        return f"Alarm set for {minutes} minutes from now: {message}"
    else:
        return "Please specify a time or number of minutes."


def send_text_message(to: str = "", message: str = "") -> str:
    """Draft a text message and read it back so user can relay it via Siri."""
    return f"Message ready for {to}: \"{message}\" — Just say: Hey Siri, text {to}, {message}."


def get_game_time(team: str = "", sport: str = "") -> str:
    """Get the next game time for a specific team."""
    sport_map = {
        "mlb": "baseball/mlb",
        "nba": "basketball/nba",
        "nhl": "hockey/nhl",
        "nfl": "football/nfl",
    }

    # Auto-detect sport from team name
    nba_teams = ["lakers", "celtics", "warriors", "nets", "cavs", "cavaliers", "bucks",
                 "suns", "heat", "knicks", "76ers", "sixers", "nuggets", "grizzlies",
                 "mavericks", "mavs", "hawks", "bulls", "raptors", "spurs", "jazz",
                 "pelicans", "kings", "magic", "pacers", "pistons", "clippers",
                 "timberwolves", "wolves", "blazers", "rockets", "thunder", "wizards", "hornets"]
    mlb_teams = ["yankees", "mets", "red sox", "dodgers", "angels", "giants", "cubs",
                 "white sox", "astros", "braves", "phillies", "padres", "mariners",
                 "twins", "guardians", "tigers", "rays", "blue jays", "orioles", "royals",
                 "rangers", "diamondbacks", "rockies", "brewers", "reds", "pirates",
                 "cardinals", "marlins", "athletics", "nationals"]
    nhl_teams = ["rangers", "islanders", "penguins", "bruins", "canadiens", "maple leafs",
                 "blackhawks", "red wings", "blues", "avalanche", "wild", "stars",
                 "predators", "lightning", "panthers", "hurricanes", "flames", "oilers",
                 "canucks", "kraken", "sharks", "ducks", "knights", "coyotes", "jets",
                 "senators", "sabres", "devils", "flyers", "capitals", "blue jackets"]

    team_lower = team.lower().strip()

    if not sport:
        if team_lower in nba_teams:
            sport = "nba"
        elif team_lower in mlb_teams:
            sport = "mlb"
        elif team_lower in nhl_teams:
            sport = "nhl"
        else:
            sport = "nba"  # default

    sport_path = sport_map.get(sport.lower(), "basketball/nba")

    try:
        resp = httpx.get(
            f"https://site.api.espn.com/apis/site/v2/sports/{sport_path}/scoreboard",
            timeout=10,
        )
        data = resp.json()
        events = data.get("events", [])

        for event in events:
            name = event.get("name", "").lower()
            short = event.get("shortName", "").lower()
            if team_lower in name or team_lower in short:
                status = event.get("status", {}).get("type", {})
                detail = status.get("shortDetail", "")
                full_name = event.get("name", "")
                date_str = event.get("date", "")

                if status.get("state") == "pre":
                    return f"{full_name} — starts at {detail}."
                elif status.get("state") == "in":
                    competitors = event.get("competitions", [{}])[0].get("competitors", [])
                    scores = []
                    for t in competitors:
                        scores.append(f"{t.get('team', {}).get('abbreviation', '?')} {t.get('score', '0')}")
                    return f"{full_name} is LIVE — {' - '.join(scores)} ({detail})."
                else:
                    competitors = event.get("competitions", [{}])[0].get("competitors", [])
                    scores = []
                    for t in competitors:
                        scores.append(f"{t.get('team', {}).get('abbreviation', '?')} {t.get('score', '0')}")
                    return f"{full_name} — Final: {' - '.join(scores)}."

        return f"No {sport.upper()} games found for {team} today."
    except Exception as e:
        return f"Could not check game time: {e}"


def homework_help(problem: str = "", subject: str = "math") -> str:
    """Solve a homework problem step by step."""
    import anthropic
    from jarvis.config import ANTHROPIC_API_KEY

    prompt = f"""You are a patient tutor helping a student. Solve this {subject} problem step by step.
Be clear and concise — explain each step in simple language.
Keep the total response under 4-5 sentences since it will be spoken aloud.
If it's a math problem, show the work briefly then give the answer.

Problem: {problem}"""

    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        resp = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=300,
            messages=[{"role": "user", "content": prompt}],
        )
        return resp.content[0].text
    except Exception as e:
        return f"Could not solve: {e}"


def summarize_url(url: str) -> str:
    """Fetch a webpage and summarize its content."""
    try:
        resp = httpx.get(url, timeout=15, follow_redirects=True,
                         headers={"User-Agent": "Mozilla/5.0"})
        html = resp.text

        # Strip HTML tags
        text = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL)
        text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)
        text = re.sub(r'<[^>]+>', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()

        if len(text) < 100:
            return "Could not extract meaningful content from that page."

        # Truncate for Claude
        text = text[:3000]

        # Use Claude to summarize
        import anthropic
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        summary_resp = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=200,
            messages=[{
                "role": "user",
                "content": f"Summarize this article in 2-3 concise sentences for spoken delivery:\n\n{text}"
            }]
        )
        return summary_resp.content[0].text
    except Exception as e:
        return f"Failed to summarize: {e}"


# Register all tools
registry.register(Tool(
    name="get_current_time",
    description="Get the current date and time.",
    parameters={"type": "object", "properties": {}, "required": []},
    handler=get_current_time,
))

registry.register(Tool(
    name="get_weather",
    description="Get weather for a city. Supports current conditions, tomorrow's forecast, or a 7-day outlook.",
    parameters={
        "type": "object",
        "properties": {
            "city": {"type": "string", "description": "City name (e.g. 'San Francisco'). Defaults to New York."},
            "when": {"type": "string", "description": "'today' for current + today's forecast, 'tomorrow' for tomorrow, 'week' for 7-day outlook. Defaults to 'today'."},
        },
        "required": [],
    },
    handler=get_weather,
))

registry.register(Tool(
    name="open_application",
    description="Open an application on the user's computer (Windows). Examples: Chrome, Notepad, VS Code, Spotify, Discord, Calculator.",
    parameters={
        "type": "object",
        "properties": {
            "name": {"type": "string", "description": "Application name to open"},
        },
        "required": ["name"],
    },
    handler=open_application,
))

registry.register(Tool(
    name="set_reminder",
    description="Save a reminder, optionally with a delay in minutes.",
    parameters={
        "type": "object",
        "properties": {
            "message": {"type": "string", "description": "The reminder message"},
            "minutes": {"type": "integer", "description": "Minutes from now (0 = no specific time)"},
        },
        "required": ["message"],
    },
    handler=set_reminder,
))

registry.register(Tool(
    name="list_reminders",
    description="List all saved reminders.",
    parameters={"type": "object", "properties": {}, "required": []},
    handler=list_reminders,
))

registry.register(Tool(
    name="web_search",
    description="Search the web for information using DuckDuckGo. Good for quick facts, definitions, and current events.",
    parameters={
        "type": "object",
        "properties": {
            "query": {"type": "string", "description": "The search query"},
        },
        "required": ["query"],
    },
    handler=web_search,
))

registry.register(Tool(
    name="run_command",
    description="Run a shell command on the user's Windows computer and return the output. Use for any system task: checking processes, network info, installing software, etc.",
    parameters={
        "type": "object",
        "properties": {
            "command": {"type": "string", "description": "The shell command to run"},
        },
        "required": ["command"],
    },
    handler=run_command,
))

registry.register(Tool(
    name="open_url",
    description="Open a URL in the default web browser.",
    parameters={
        "type": "object",
        "properties": {
            "url": {"type": "string", "description": "The URL to open"},
        },
        "required": ["url"],
    },
    handler=open_url,
))

registry.register(Tool(
    name="read_file",
    description="Read the contents of a file on the computer.",
    parameters={
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Full file path to read"},
        },
        "required": ["path"],
    },
    handler=read_file,
))

registry.register(Tool(
    name="write_file",
    description="Write content to a file. Creates the file if it doesn't exist.",
    parameters={
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Full file path to write"},
            "content": {"type": "string", "description": "Content to write"},
        },
        "required": ["path", "content"],
    },
    handler=write_file,
))

registry.register(Tool(
    name="list_directory",
    description="List files and folders in a directory.",
    parameters={
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Directory path (default: current directory)"},
        },
        "required": [],
    },
    handler=list_directory,
))

registry.register(Tool(
    name="kill_process",
    description="Kill a running process by name (e.g. 'chrome.exe', 'notepad.exe').",
    parameters={
        "type": "object",
        "properties": {
            "name": {"type": "string", "description": "Process name to kill (e.g. 'chrome.exe')"},
        },
        "required": ["name"],
    },
    handler=kill_process,
))

registry.register(Tool(
    name="get_system_info",
    description="Get system information — CPU, RAM usage, disk space, battery level.",
    parameters={"type": "object", "properties": {}, "required": []},
    handler=get_system_info,
))

registry.register(Tool(
    name="screenshot",
    description="Take a screenshot of the screen and save it to the desktop.",
    parameters={"type": "object", "properties": {}, "required": []},
    handler=screenshot,
))

registry.register(Tool(
    name="play_music",
    description="Play music by searching Spotify or YouTube. Use for requests like 'play Rodeo', 'put on Drake', 'play MBDTF on YouTube'.",
    parameters={
        "type": "object",
        "properties": {
            "query": {"type": "string", "description": "Song, artist, album, or playlist to search for"},
            "service": {"type": "string", "description": "'spotify' (default) or 'youtube'"},
        },
        "required": ["query"],
    },
    handler=play_music,
))

registry.register(Tool(
    name="control_music",
    description="Control media playback — play, pause, next/skip, previous/back, volume_up, volume_down, mute.",
    parameters={
        "type": "object",
        "properties": {
            "action": {"type": "string", "description": "Action: play, pause, next, skip, previous, back, volume_up, volume_down, mute"},
        },
        "required": ["action"],
    },
    handler=control_music,
))

registry.register(Tool(
    name="get_news",
    description="Get latest news headlines. Can filter by topic (e.g. 'sports', 'tech', 'Kalshi', 'stock market'). Use when user asks 'what's in the news?' or 'any news about...'.",
    parameters={
        "type": "object",
        "properties": {
            "topic": {"type": "string", "description": "News topic to search for (empty for top headlines)"},
        },
    },
    handler=get_news,
))

registry.register(Tool(
    name="create_document",
    description="Create a Word document (.docx) or text file (.txt) on the Desktop. Use when user says 'make a document', 'write a doc', 'create a file about...'. Jarvis generates the content and saves it.",
    parameters={
        "type": "object",
        "properties": {
            "title": {"type": "string", "description": "Document title / filename"},
            "content": {"type": "string", "description": "Document body text. Use # for headings, ## for subheadings, - for bullet points, newlines for paragraphs."},
            "format": {"type": "string", "description": "'docx' (default) or 'txt'"},
        },
        "required": ["title", "content"],
    },
    handler=create_document,
))

registry.register(Tool(
    name="draft_email",
    description="Compose and open an email draft in the user's default email app. Use when user says 'write an email to...' or 'email...'.",
    parameters={
        "type": "object",
        "properties": {
            "to": {"type": "string", "description": "Recipient email address"},
            "subject": {"type": "string", "description": "Email subject line"},
            "body": {"type": "string", "description": "Email body text"},
        },
        "required": ["to", "subject", "body"],
    },
    handler=draft_email,
))

registry.register(Tool(
    name="analyze_screenshot",
    description="Take a screenshot and analyze what's on the screen using AI vision. Use when user asks 'what's on my screen?', 'what am I looking at?', or 'analyze my screen'.",
    parameters={"type": "object", "properties": {}},
    handler=analyze_screenshot,
))

registry.register(Tool(
    name="set_alarm",
    description="Set an alarm — either at a specific time ('7:00 AM') or in X minutes ('10 minutes'). Use when user says 'set alarm for...', 'wake me up at...', 'alarm in 10 minutes'.",
    parameters={
        "type": "object",
        "properties": {
            "minutes": {"type": "integer", "description": "Minutes from now (use this OR time_str)"},
            "time_str": {"type": "string", "description": "Specific time like '7:00 AM' or '19:00' (use this OR minutes)"},
            "message": {"type": "string", "description": "Alarm message/label (default: 'Alarm')"},
        },
        "required": [],
    },
    handler=set_alarm,
))

registry.register(Tool(
    name="send_notification",
    description="Send a push notification to Deagz's phone via ntfy app. Use for 'send me a notification', 'message my phone', 'notify me', 'send to my phone', 'ntfy'.",
    parameters={
        "type": "object",
        "properties": {
            "message": {"type": "string", "description": "The notification message"},
            "title": {"type": "string", "description": "Notification title (default: JARVIS)"},
        },
        "required": ["message"],
    },
    handler=send_notification,
))

registry.register(Tool(
    name="send_text",
    description="Send a text message to someone. Opens the SMS app with the message ready to send. Use when user says 'text mom', 'send a message to...', 'text [contact]'.",
    parameters={
        "type": "object",
        "properties": {
            "to": {"type": "string", "description": "Phone number or contact name"},
            "message": {"type": "string", "description": "The message to send"},
        },
        "required": ["to", "message"],
    },
    handler=send_text_message,
))

registry.register(Tool(
    name="get_game_time",
    description="Get the next game time, live score, or final score for a specific team. Use when user asks 'when do the Cavs play?', 'what time is the Lakers game?', 'did the Cubs win?'.",
    parameters={
        "type": "object",
        "properties": {
            "team": {"type": "string", "description": "Team name (e.g. 'Cavs', 'Lakers', 'Yankees')"},
            "sport": {"type": "string", "description": "Sport: nba, mlb, nhl, nfl (auto-detected if omitted)"},
        },
        "required": ["team"],
    },
    handler=get_game_time,
))

registry.register(Tool(
    name="homework_help",
    description="Solve a homework problem step by step. Works for math, algebra, science, history, etc. Use when user says 'help me with this problem', 'solve...', 'what is...', or pastes a homework question.",
    parameters={
        "type": "object",
        "properties": {
            "problem": {"type": "string", "description": "The homework problem to solve"},
            "subject": {"type": "string", "description": "Subject area: math, algebra, science, history, etc. (default: math)"},
        },
        "required": ["problem"],
    },
    handler=homework_help,
))

registry.register(Tool(
    name="summarize_url",
    description="Fetch a webpage/article URL and summarize its content in 2-3 sentences. Use when user says 'summarize this article' or 'read this for me'.",
    parameters={
        "type": "object",
        "properties": {
            "url": {"type": "string", "description": "The full URL to fetch and summarize"},
        },
        "required": ["url"],
    },
    handler=summarize_url,
))


# ─── PC Power Management ───

def lock_computer(**kwargs) -> str:
    """Lock the computer screen."""
    try:
        subprocess.run("rundll32.exe user32.dll,LockWorkStation", shell=True, timeout=5)
        return "Computer locked."
    except Exception as e:
        return f"Failed to lock: {e}"


def sleep_computer(delay_minutes: int = 0, **kwargs) -> str:
    """Put the computer to sleep, optionally after a delay."""
    import threading

    def _sleep():
        subprocess.run(
            "rundll32.exe powrprof.dll,SetSuspendState 0,1,0",
            shell=True, timeout=10,
        )

    if delay_minutes > 0:
        threading.Timer(delay_minutes * 60, _sleep).start()
        return f"Computer will sleep in {delay_minutes} minutes."
    else:
        _sleep()
        return "Putting computer to sleep."


def shutdown_computer(delay_minutes: int = 0, **kwargs) -> str:
    """Schedule a computer shutdown. Default: immediate. Use delay_minutes for timed shutdown."""
    try:
        if delay_minutes > 0:
            seconds = delay_minutes * 60
            subprocess.run(
                f"shutdown /s /t {seconds} /c \"JARVIS scheduled shutdown\"",
                shell=True, timeout=5,
            )
            return f"Shutdown scheduled in {delay_minutes} minutes. Say 'cancel shutdown' to abort."
        else:
            subprocess.run("shutdown /s /t 30 /c \"JARVIS shutdown\"", shell=True, timeout=5)
            return "Shutting down in 30 seconds. Say 'cancel shutdown' to abort."
    except Exception as e:
        return f"Failed to schedule shutdown: {e}"


def restart_computer(delay_minutes: int = 0, **kwargs) -> str:
    """Restart the computer, optionally after a delay."""
    try:
        if delay_minutes > 0:
            seconds = delay_minutes * 60
            subprocess.run(
                f"shutdown /r /t {seconds} /c \"JARVIS scheduled restart\"",
                shell=True, timeout=5,
            )
            return f"Restart scheduled in {delay_minutes} minutes. Say 'cancel shutdown' to abort."
        else:
            subprocess.run("shutdown /r /t 30 /c \"JARVIS restart\"", shell=True, timeout=5)
            return "Restarting in 30 seconds. Say 'cancel shutdown' to abort."
    except Exception as e:
        return f"Failed to schedule restart: {e}"


def cancel_shutdown(**kwargs) -> str:
    """Cancel a scheduled shutdown or restart."""
    try:
        subprocess.run("shutdown /a", shell=True, timeout=5)
        return "Scheduled shutdown/restart cancelled."
    except Exception as e:
        return f"Failed to cancel: {e}"


def set_brightness(level: int = 50, **kwargs) -> str:
    """Set screen brightness (0-100)."""
    try:
        level = max(0, min(100, level))
        subprocess.run(
            f'powershell -Command "(Get-WmiObject -Namespace root/WMI '
            f'-Class WmiMonitorBrightnessMethods).WmiSetBrightness(1,{level})"',
            shell=True, capture_output=True, timeout=10,
        )
        return f"Brightness set to {level}%."
    except Exception as e:
        return f"Failed to set brightness: {e}"


# Register power management tools
registry.register(Tool(
    name="lock_computer",
    description="Lock the computer screen. Use when user says 'lock my computer', 'lock the screen', 'lock it'.",
    parameters={"type": "object", "properties": {}},
    handler=lock_computer,
))

registry.register(Tool(
    name="sleep_computer",
    description="Put the computer to sleep. Optionally delay by X minutes. Use for 'put my computer to sleep', 'sleep in 30 minutes'.",
    parameters={
        "type": "object",
        "properties": {
            "delay_minutes": {
                "type": "integer",
                "description": "Minutes to wait before sleeping (0 = immediate)",
            },
        },
    },
    handler=sleep_computer,
))

registry.register(Tool(
    name="shutdown_computer",
    description="Shut down the computer. Optionally delay by X minutes. Use for 'shut down', 'turn off my computer', 'shutdown in 1 hour'.",
    parameters={
        "type": "object",
        "properties": {
            "delay_minutes": {
                "type": "integer",
                "description": "Minutes to wait before shutdown (0 = 30 second countdown)",
            },
        },
    },
    handler=shutdown_computer,
))

registry.register(Tool(
    name="restart_computer",
    description="Restart the computer. Optionally delay by X minutes. Use for 'restart', 'reboot my computer'.",
    parameters={
        "type": "object",
        "properties": {
            "delay_minutes": {
                "type": "integer",
                "description": "Minutes to wait before restart (0 = 30 second countdown)",
            },
        },
    },
    handler=restart_computer,
))

registry.register(Tool(
    name="cancel_shutdown",
    description="Cancel a scheduled shutdown or restart. Use when user says 'cancel shutdown', 'don't shut down', 'abort restart'.",
    parameters={"type": "object", "properties": {}},
    handler=cancel_shutdown,
))

registry.register(Tool(
    name="set_brightness",
    description="Set screen brightness (0-100). Use for 'set brightness to 50', 'dim the screen', 'turn brightness up'.",
    parameters={
        "type": "object",
        "properties": {
            "level": {
                "type": "integer",
                "description": "Brightness level 0-100",
            },
        },
        "required": ["level"],
    },
    handler=set_brightness,
))

registry.register(Tool(
    name="set_volume",
    description="Set system volume (0-100). Use for 'set volume to 50', 'turn it down', 'volume up'.",
    parameters={
        "type": "object",
        "properties": {
            "level": {
                "type": "integer",
                "description": "Volume level 0-100",
            },
        },
        "required": ["level"],
    },
    handler=set_volume,
))
