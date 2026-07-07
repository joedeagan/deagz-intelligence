"""FastAPI server for JARVIS web interface."""

import sys
import os
import time as _tz_time

# The cloud host runs UTC — anchor the whole process to home time so every
# datetime.now() across the brain's tools reports Akron time.
# LINUX ONLY: on Windows the CRT misparses IANA names like America/New_York
# and skews the clock hours off — the laptop's own local timezone is already
# correct, so leave it alone there.
if os.name != "nt":
    os.environ.setdefault("TZ", "America/New_York")
    try:
        _tz_time.tzset()
    except AttributeError:
        pass

# Add project root to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from fastapi import FastAPI, UploadFile, File
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, Response, StreamingResponse
from pydantic import BaseModel
import httpx
import asyncio
import edge_tts
import soundfile as sf
import numpy as np
import tempfile
from pathlib import Path

from jarvis.config import (
    ELEVENLABS_API_KEY,
    ELEVENLABS_VOICE_ID,
    FISH_AUDIO_API_KEY,
    FISH_AUDIO_MODEL_ID,
    JARVIS_VOICE,
    TTS_ENGINE,
)
from jarvis.tools.voice import get_active_voice

# Import tools to register them
from jarvis.tools import system as _s  # noqa
from jarvis.tools import kalshi as _k  # noqa
from jarvis.tools import study as _st  # noqa
from jarvis.tools import shazam as _sh  # noqa
from jarvis.tools import spotify as _sp2  # noqa
from jarvis.tools import memory as _mem  # noqa
from jarvis.tools import image_gen as _img  # noqa
from jarvis.tools import voice as _vc  # noqa
from jarvis.tools import kalshi_advisor as _ka  # noqa
from jarvis.tools import autodj as _dj  # noqa
from jarvis.tools import screen_aware as _sa  # noqa
from jarvis.tools import proactive as _pro  # noqa
from jarvis.tools import coder as _code  # noqa
from jarvis.tools import contacts as _ct  # noqa
from jarvis.tools import routines as _rt  # noqa
from jarvis.tools import backtester as _bt  # noqa
from jarvis.tools import stems as _stems  # noqa
from jarvis.tools import sports as _spt  # noqa
from jarvis.tools import selfbuild as _sb  # noqa
from jarvis.brain import Brain

app = FastAPI(title="JARVIS")
brain = Brain()

# Serve static files
static_dir = os.path.join(os.path.dirname(__file__), "static")
app.mount("/static", StaticFiles(directory=static_dir), name="static")


class ChatRequest(BaseModel):
    message: str
    context: str = ""  # the wall's live state — art, timers, what's playing


class IntentRequest(BaseModel):
    text: str
    movies: list = []


class TTSRequest(BaseModel):
    text: str


def fix_pronunciation(text: str) -> str:
    for v in ("Deagz", "deagz", "DEAGZ"):
        text = text.replace(v, "Deegz")
    return text


@app.get("/")
async def index():
    return FileResponse(os.path.join(static_dir, "index.html"))

@app.get("/wall")
async def wall():
    """iPad wall display — flip clock with built-in Jarvis voice."""
    return FileResponse(
        os.path.join(static_dir, "wall.html"),
        headers={"Cache-Control": "no-store, must-revalidate"},  # Safari silently served stale walls for hours
    )


# === Home agent command queue ===
# The wall page enqueues commands; the agent on the home laptop polls for
# them (outbound-only from the house) and executes them against Jellyfin/TV.
import time as _time

_agent_queue: list = []
AGENT_CMD_TTL = 90  # seconds before an unclaimed command goes stale

class AgentCommand(BaseModel):
    type: str
    payload: dict = {}

@app.post("/api/agent/enqueue")
async def agent_enqueue(cmd: AgentCommand):
    _agent_queue.append({"type": cmd.type, "payload": cmd.payload, "ts": _time.time()})
    return {"queued": True}

@app.get("/api/agent/poll")
async def agent_poll():
    """Home agent calls this every few seconds; hands over pending commands."""
    global _agent_queue
    now = _time.time()
    fresh = [c for c in _agent_queue if now - c["ts"] < AGENT_CMD_TTL]
    _agent_queue = []
    return {"commands": fresh}


# Separate queue for the desktop PC listener (so it and the home agent don't
# steal each other's commands off one shared queue).
_pc_queue: list = []

@app.post("/api/pc/enqueue")
async def pc_enqueue(cmd: AgentCommand):
    _pc_queue.append({"type": cmd.type, "payload": cmd.payload, "ts": _time.time()})
    return {"queued": True}

@app.get("/api/pc/poll")
async def pc_poll():
    global _pc_queue
    now = _time.time()
    fresh = [c for c in _pc_queue if now - c["ts"] < AGENT_CMD_TTL]
    _pc_queue = []
    return {"commands": fresh}


# === Live movie library ===
@app.get("/api/library")
def library():
    """The wall's movie list, live from Jellyfin on this same laptop.

    New files become voice-playable the moment Jellyfin indexes them.
    Falls back to the static library.json on the cloud instance (which
    can't reach the house) or if Jellyfin is napping.
    """
    import json as _json
    from pathlib import Path as _P

    key = os.getenv("JELLYFIN_API_KEY", "")
    if not key:
        try:  # the home agent's config on this same machine already holds the key
            key = _json.loads(_P("C:/jarvis-agent/config.json").read_text()).get("api_key", "")
        except Exception:
            key = ""
    try:
        r = httpx.get(
            "http://127.0.0.1:8096/Items",
            params={"IncludeItemTypes": "Movie", "Recursive": "true",
                    "fields": "ProductionYear", "api_key": key},
            timeout=8,
        )
        items = r.json().get("Items", [])
        sid = httpx.get("http://127.0.0.1:8096/System/Info/Public", timeout=5).json().get("Id", "")
        return {
            "host": "192.168.1.73", "port": 3010, "serverId": sid, "live": True,
            "movies": [{"name": i.get("Name"), "year": i.get("ProductionYear"), "id": i.get("Id")}
                       for i in items],
        }
    except Exception:
        try:
            return _json.loads(open(os.path.join(static_dir, "library.json"), encoding="utf-8").read())
        except Exception:
            return {"host": "192.168.1.73", "port": 3010, "movies": [], "live": False}


# === Duck media while the user talks (auto pause/resume) ===
def _jellyfin_key():
    import json as _json
    from pathlib import Path as _P
    key = os.getenv("JELLYFIN_API_KEY", "")
    if not key:
        try:
            key = _json.loads(_P("C:/jarvis-agent/config.json").read_text()).get("api_key", "")
        except Exception:
            key = ""
    return key


def _tv_is_playing():
    try:
        key = _jellyfin_key()
        r = httpx.get("http://127.0.0.1:8096/Sessions", params={"api_key": key}, timeout=5)
        for s in r.json():
            if "WebOS" in (s.get("Client") or "") or "LG" in (s.get("DeviceName") or ""):
                if s.get("NowPlayingItem") and not (s.get("PlayState") or {}).get("IsPaused", False):
                    return True
    except Exception:
        pass
    return False


def _spotify_is_playing():
    try:
        from jarvis.tools.spotify import _get_spotify
        sp = _get_spotify()
        cur = sp.current_playback() if sp else None
        return bool(cur and cur.get("is_playing"))
    except Exception:
        return False


@app.post("/api/media/duck")
async def media_duck():
    """Pause whatever is actively playing (TV / Spotify) so the mic can hear
    the user. Returns which sources were paused, to resume exactly those."""
    ducked = {"tv": False, "spotify": False}
    if _spotify_is_playing():
        try:
            from jarvis.tools.spotify import _get_spotify
            _get_spotify().pause_playback()
            ducked["spotify"] = True
        except Exception:
            pass
    if _tv_is_playing():
        _agent_queue.append({"type": "tv_command", "payload": {"command": "Pause"}, "ts": _time.time()})
        ducked["tv"] = True
    return ducked


class UnduckRequest(BaseModel):
    tv: bool = False
    spotify: bool = False


@app.post("/api/media/unduck")
async def media_unduck(req: UnduckRequest):
    if req.spotify:
        try:
            from jarvis.tools.spotify import _get_spotify
            _get_spotify().start_playback()
        except Exception:
            pass
    if req.tv:
        _agent_queue.append({"type": "tv_command", "payload": {"command": "Unpause"}, "ts": _time.time()})
    return {"ok": True}


# === House intercom ===
# Anyone posts a line; the wall polls and speaks it. On the cloud instance
# the home agent relays announcements down to the local brain.
_announcements: list = []

class AnnounceRequest(BaseModel):
    text: str

@app.post("/api/announce")
async def announce(req: AnnounceRequest):
    text = (req.text or "").strip()[:300]
    if text:
        _announcements.append({"text": text, "ts": _time.time()})
    return {"ok": bool(text)}

@app.get("/api/announcements")
async def announcements():
    global _announcements
    now = _time.time()
    fresh = [a for a in _announcements if now - a["ts"] < 120]
    _announcements = []
    return {"announcements": fresh}

# === Sleepy History (bedtime YouTube narrator) ===
# YouTube killed public channel RSS, so the videos tab gets scraped instead
# (cached 6h). The wall's embedded player plays the pick without ever
# leaving the dashboard.
SLEEPY_CHANNEL = os.getenv("SLEEPY_CHANNEL", "@HistorianSleepy")
_sleepy_cache = {"ids": [], "ts": 0.0}


@app.get("/api/sleepy")
def sleepy(pick: str = "latest"):
    import re as _re
    import random as _random
    now = _time.time()
    if now - _sleepy_cache["ts"] > 6 * 3600 or not _sleepy_cache["ids"]:
        try:
            r = httpx.get(
                f"https://www.youtube.com/{SLEEPY_CHANNEL}/videos",
                headers={"User-Agent": "Mozilla/5.0", "Accept-Language": "en-US"},
                timeout=15,
            )
            ids = []
            for vid in _re.findall(r'"videoId":"([\w-]{11})"', r.text):
                if vid not in ids:
                    ids.append(vid)
            if ids:
                _sleepy_cache["ids"] = ids[:20]
                _sleepy_cache["ts"] = now
        except Exception:
            pass  # keep last known list
    ids = _sleepy_cache["ids"]
    if not ids:
        return {"videoId": "", "title": ""}
    vid = ids[0] if pick == "latest" else _random.choice(ids)
    title = ""
    try:  # official oEmbed still works — one call per pick for the spoken title
        o = httpx.get(
            "https://www.youtube.com/oembed",
            params={"url": f"https://www.youtube.com/watch?v={vid}", "format": "json"},
            timeout=10,
        )
        title = o.json().get("title", "")
    except Exception:
        pass
    return {"videoId": vid, "title": title}


# === iMessage intercom ===
# The wall iPad's Shortcuts automation fires on incoming texts (sender-
# filtered) and POSTs them here. The wall speaks the message like an intercom
# line, and the response carries a Jarvis-composed reply the Shortcut sends
# back as a real iMessage.
class IMessageIn(BaseModel):
    text: str
    sender: str = "the family"


@app.post("/api/imessage")
def imessage_in(req: IMessageIn):
    text = (req.text or "").strip()[:500]
    if not text:
        return {"reply": ""}
    sender = (req.sender or "the family").strip()[:60]
    _announcements.append({"text": f"Text from {sender}, sir: {text}", "ts": _time.time()})

    reply = "got it"
    try:
        import anthropic
        from jarvis.config import ANTHROPIC_API_KEY
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        msg = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=60,
            system=(
                "You are auto-replying to a family text FROM Joe Deagan's own phone "
                "number, so you must sound exactly like Joe texting. ONE short reply "
                "- lowercase, brief, minimal punctuation, never formal, never "
                "assistant-like. MATCH THE TONE TO WHO TEXTED: grandparents get warm "
                "and sweet ('hi grandma', 'ok love you', 'sounds good, thanks "
                "grandma') - NO slang like 'yo', 'whats up', 'bruh' with them; "
                "parents/siblings get plain casual ('ok', 'got it', 'omw'). Just "
                "acknowledge the message. HARD RULES: never commit Joe to plans, "
                "never say where he is or what he's doing - if they ask something "
                "only Joe can answer, reply like 'ill text u back in a sec'. The "
                "message was also announced on his room display, so he genuinely "
                "will see it."
            ),
            messages=[{"role": "user", "content": f"{sender} texted: {text}"}],
        )
        composed = " ".join(b.text for b in msg.content if b.type == "text").strip()
        if composed:
            reply = composed
    except Exception:
        pass  # the stock line above still goes out
    return {"reply": reply}


# The observer: proactive Jarvis. New movies, incoming rain, Kalshi moves —
# announced through the same queue the intercom uses, so the wall just speaks
# them. HOUSE BRAIN ONLY (Windows laptop): the cloud fallback runs this same
# file on Linux and must not double-speak into the room.
@app.on_event("startup")
def _start_observer():
    if os.name == "nt":
        from jarvis.tools.observer import start_observer
        from jarvis.tools.mind import start_mind

        def _observer_announce(text: str):
            _announcements.append({"text": str(text)[:300], "ts": _time.time()})

        start_observer(_observer_announce)
        start_mind(_observer_announce)  # the inner life — he thinks hourly, speaks rarely
        _sb.set_announcer(_observer_announce)  # self-build drafts announce themselves
    _sb.load_selfbuilt()  # bring his self-built abilities online (safe no-op elsewhere)


@app.get("/announce")
async def announce_page():
    return FileResponse(
        os.path.join(static_dir, "announce.html"),
        headers={"Cache-Control": "no-store"},
    )

@app.get("/sw.js")
async def service_worker():
    return FileResponse(os.path.join(static_dir, "sw.js"), media_type="application/javascript")



@app.get("/api/dashboard")
async def dashboard():
    """Aggregated dashboard data — weather, portfolio, bot status."""
    import datetime
    from jarvis.tools.system import get_weather
    from jarvis.tools.kalshi import get_portfolio, get_bot_status

    data = {
        "time": datetime.datetime.now().strftime("%I:%M %p"),
        "date": datetime.datetime.now().strftime("%A, %B %d"),
        "weather": {},
        "portfolio": {},
        "bot": {},
    }

    # Weather
    try:
        weather_raw = get_weather()
        lines = weather_raw.split("\n")
        data["weather"]["summary"] = lines[0] if lines else "N/A"
        data["weather"]["forecast"] = lines[1] if len(lines) > 1 else ""
    except Exception:
        data["weather"]["summary"] = "Weather unavailable"

    # Portfolio
    try:
        port = _get_portfolio_data()
        data["portfolio"] = port
    except Exception:
        data["portfolio"] = {"balance": 0, "pnl": 0, "positions": 0}

    # Bot status
    try:
        from jarvis.tools.kalshi import _get
        bot = _get("/api/bot/status")
        if isinstance(bot, dict):
            data["bot"] = {
                "running": bot.get("running", False),
                "scans": bot.get("scan_count", 0),
                "trades": bot.get("trades_today", bot.get("exits_today", 0)),
                "pnl": bot.get("daily_loss_cents", 0) / 100,
            }
    except Exception:
        data["bot"] = {"running": False}

    # Equity chart data
    try:
        from jarvis.tools.kalshi import _get as kalshi_get
        eq = kalshi_get("/api/bot/equity?days=7")
        if isinstance(eq, dict):
            points = eq.get("equity", [])
            data["equity_chart"] = [p.get("equity_cents", 0) / 100 for p in points[-30:]]
    except Exception:
        data["equity_chart"] = []

    return data


def _get_portfolio_data():
    """Get portfolio data as a dict for the dashboard."""
    from jarvis.tools.kalshi import _get
    raw = _get("/api/portfolio")
    if isinstance(raw, str):
        return {"balance": 0, "pnl": 0, "positions": 0}

    balance = raw.get("balance", raw.get("balance_cents", 0))
    if balance > 1000:
        balance /= 100

    return {
        "balance": round(balance, 2),
        "portfolio_value": round(raw.get("portfolio_value", 0), 2),
        "pnl": round(raw.get("unrealised_pnl", 0), 2),
        "positions": raw.get("position_count", len(raw.get("positions", []))),
    }


# === Movie companion ===
# The TV's Jellyfin app reports what's playing and the EXACT position — the
# same data media-ducking reads. Every chat question gets it as context with
# a hard no-spoilers rule, so "wait, who is that guy?" mid-movie gets answered
# as of YOUR minute, never past it.
def _tv_movie_context() -> str:
    try:
        r = httpx.get("http://127.0.0.1:8096/Sessions",
                      params={"api_key": _jellyfin_key()}, timeout=4)
        for s in r.json():
            if "WebOS" in (s.get("Client") or "") or "LG" in (s.get("DeviceName") or ""):
                item = s.get("NowPlayingItem")
                if not item:
                    continue
                state = s.get("PlayState") or {}
                pos_min = int(state.get("PositionTicks", 0) // 600_000_000)  # ticks -> minutes
                total_min = int((item.get("RunTimeTicks") or 0) // 600_000_000)
                name = item.get("Name", "")
                year = item.get("ProductionYear", "")
                verb = "paused at" if state.get("IsPaused") else "playing, currently at"
                return (f"The TV is watching {name} ({year}) — {verb} minute {pos_min}"
                        + (f" of {total_min}" if total_min else "") + ". "
                        f"SPOILER RULE (absolute): if he asks about this film, answer only "
                        f"with what a first-time viewer knows at minute {pos_min} — who a "
                        f"character is, what just happened. NEVER reveal, foreshadow, or "
                        f"hint at ANYTHING after that minute, even if asked directly; "
                        f"offer to say more once he's watched further.")
    except Exception:
        pass
    return ""


def _with_tv_context(context: str) -> str:
    parts = [p for p in (context, _tv_movie_context()) if p]
    try:  # live TV-app + PC status, reported by the house agents
        from jarvis.tools import housestate
        hs = housestate.snapshot()
        if hs:
            parts.append(hs)
    except Exception:
        pass
    return " | ".join(parts)


class DeviceReport(BaseModel):
    device: str  # "tv" | "pc"
    info: dict = {}


@app.post("/api/housestate")
def housestate_report(req: DeviceReport):
    from jarvis.tools import housestate
    housestate.report(req.device, req.info)
    return {"ok": True}


@app.post("/api/chat")
async def chat(req: ChatRequest):
    response = brain.think(req.message, context=_with_tv_context(req.context))
    return {"response": response}


import hashlib

# In-memory TTS cache + disk cache
_tts_cache: dict[str, bytes] = {}
_cache_dir = Path(os.path.dirname(__file__)).parent / "data" / "tts_cache"
_cache_dir.mkdir(parents=True, exist_ok=True)


def _cache_key(text: str) -> str:
    return hashlib.md5(text.lower().strip().encode()).hexdigest()


def _load_cache():
    """Load cached audio files from disk on startup."""
    for f in _cache_dir.glob("*.mp3"):
        _tts_cache[f.stem] = f.read_bytes()
    print(f"TTS cache: {len(_tts_cache)} entries loaded from disk")


def _truncate_for_speech(text: str, max_chars: int = 1000) -> str:
    """Light truncation — only cuts extremely long responses."""
    if len(text) <= max_chars:
        return text
    cutoff = text[:max_chars]
    for end in [". ", "! ", "? "]:
        idx = cutoff.rfind(end)
        if idx > 100:
            return cutoff[:idx + 1]
    return cutoff.rstrip() + "."


# Load disk cache on import
_load_cache()

# Persistent HTTP client for ElevenLabs (avoids connection setup per request)
_eleven_client = httpx.AsyncClient(timeout=30)


async def generate_tts(text: str) -> bytes:
    """Generate TTS audio bytes from text. Uses cache for repeated phrases."""
    text = fix_pronunciation(text)
    if not text:
        return b""

    # Truncate for speech to save credits
    text = _truncate_for_speech(text)

    # Check cache first — instant response
    key = _cache_key(text)
    if key in _tts_cache:
        return _tts_cache[key]

    # Try Fish Audio first
    if TTS_ENGINE == "fish" and FISH_AUDIO_API_KEY:
        try:
            async with httpx.AsyncClient() as client:
                resp = await client.post(
                    "https://api.fish.audio/v1/tts",
                    headers={
                        "Authorization": f"Bearer {FISH_AUDIO_API_KEY}",
                        "Content-Type": "application/json",
                    },
                    json={
                        "text": text,
                        "reference_id": FISH_AUDIO_MODEL_ID,
                        "format": "mp3",
                        "mp3_bitrate": 128,
                    },
                    timeout=30,
                )
            if resp.status_code == 200:
                _save_to_cache(text, resp.content)
                return resp.content
            else:
                print(f"Fish Audio error: {resp.status_code} {resp.text}")
        except Exception as e:
            print(f"Fish Audio exception: {e}")

    # Try ElevenLabs with flash model + persistent client (fastest)
    if (TTS_ENGINE == "elevenlabs" or TTS_ENGINE == "fish") and ELEVENLABS_API_KEY:
        try:
            voice = get_active_voice()
            voice_id = voice.get("voice_id", ELEVENLABS_VOICE_ID)
            resp = await _eleven_client.post(
                f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}",
                headers={
                    "xi-api-key": ELEVENLABS_API_KEY,
                    "Content-Type": "application/json",
                },
                json={
                    "text": text,
                    "model_id": "eleven_flash_v2_5",
                    "voice_settings": {
                        "stability": 0.75,
                        "similarity_boost": 0.80,
                    },
                    "optimize_streaming_latency": 4,
                },
            )
            if resp.status_code == 200:
                _save_to_cache(text, resp.content)
                return resp.content
        except Exception:
            pass

    # No fallback to Edge TTS — user hates it. Return empty audio (text still shows on screen)
    print("[TTS] ElevenLabs failed or out of credits. No voice — text only.")
    return b""


def _save_to_cache(text: str, audio: bytes):
    """Save TTS audio to both memory and disk cache."""
    if len(audio) < 100:  # skip empty/broken audio
        return
    key = _cache_key(text)
    _tts_cache[key] = audio
    cache_file = _cache_dir / f"{key}.mp3"
    cache_file.write_bytes(audio)
    # Keep cache under 200 entries
    if len(_tts_cache) > 200:
        oldest = next(iter(_tts_cache))
        del _tts_cache[oldest]
        (cache_dir / f"{oldest}.mp3").unlink(missing_ok=True)


@app.post("/api/tts")
async def tts(req: TTSRequest):
    audio = await generate_tts(req.text)
    if not audio:
        return Response(status_code=400)
    return Response(content=audio, media_type="audio/mpeg")


@app.post("/api/intent")
def intent(req: IntentRequest):
    """Natural-language command router for the wall.

    Fast model maps ANY phrasing ('turn tv volume to fifteen') onto one of
    the wall's device commands; 'none' falls through to normal chat.
    Runs only when the wall's quick regex patterns miss.
    """
    import json as _json
    import re as _re
    import anthropic
    from jarvis.config import ANTHROPIC_API_KEY

    movie_names = ", ".join(str(m) for m in req.movies[:60]) or "(none)"
    system = (
        "Map the user's utterance to ONE home command. Movies available: " + movie_names + ".\n"
        'Reply ONLY with JSON like {"intent": "..."} plus needed fields.\n'
        "Intents: volume_set(level 0-100), volume_up, volume_down, mute, unmute, "
        "tv_off, tv_on, open_app(app: netflix|youtube|youtube tv|jellyfin|prime|disney|hulu|spotify|<any app name>), "
        "play_movie(title, on_tv: true/false), play_music, pause, resume, stop_playback, "
        "timer(seconds), alarm(hour 0-23, minute 0-59), movie_list, tv_message(text), "
        "look (asking what you can see / to look at something via camera), "
        "briefing (morning report / summary of the day), music_control, "
        "goodnight (goodnight / going to bed / going to sleep = wind the room down), "
        "good_morning (a good-morning greeting = wake the wall + morning briefing), "
        "sleepy_history(minutes, pick: latest|random) = play a Historian Sleepy bedtime "
        "story/narration on the wall (default minutes 60; 'a random one' = pick random), "
        "sleep_sounds(kind: rain|white|brown, minutes) = ambient sleep noise on the wall, "
        "stop_sounds = stop the bedtime story or sleep sounds, "
        "pause_story = pause the bedtime story, resume_story = continue it, "
        "paint_wall(prompt) = generate/change the wall's backdrop image "
        "(prompt is the scene description; empty prompt = clear it; QUESTIONS "
        "about the wall/backdrop like 'what did you paint' = none, never paint_wall), "
        "pc(action: on|lock|sleep|shutdown|restart, app) = control the desktop PC "
        "(action 'on' wakes it; app set only when launching something; GAMES like "
        "fortnite/minecraft/roblox always mean pc with that app), none.\n"
        "play_music = ANY request to play songs/artists/albums/playlists/music (Spotify etc). "
        "music_control = pause/resume/skip/previous/volume changes when it is about MUSIC or Spotify "
        "(pause/resume/stop_playback and volume_set/up/down are for the TV/movies only). "
        "open_app ONLY when they explicitly say open/launch an app by name. "
        "play_movie ONLY for titles in the movie list.\n"
        "Use none for questions, conversation, or anything that is not a device command."
    )
    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        msg = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=150,
            system=system,
            messages=[{"role": "user", "content": req.text}],
        )
        raw = msg.content[0].text
        m = _re.search(r"\{.*\}", raw, _re.S)
        return _json.loads(m.group(0)) if m else {"intent": "none"}
    except Exception:
        return {"intent": "none"}


@app.get("/api/pvkey")
def pvkey():
    """Picovoice access key for the on-device 'Jarvis' wake word (from .env)."""
    return {"key": os.getenv("PICOVOICE_ACCESS_KEY", "")}


# --- local wake-word gate (openWakeWord, "Hey Jarvis") -----------------------
# The ear posts every utterance here BEFORE paying for transcription. The
# check runs entirely on this laptop: free, fast, and it kills background-TV
# false triggers. If the engine is missing, we answer wake=True so the wall
# falls back to the old transcribe-everything behavior.
_oww = None
_oww_lock = None


@app.post("/api/wake")
def wake_check(audio: UploadFile = File(...)):
    global _oww, _oww_lock
    import threading
    if _oww_lock is None:
        _oww_lock = threading.Lock()
    data = audio.file.read()
    result = {"ok": False, "wake": True}
    try:
        with _oww_lock:
            if _oww is None:
                from openwakeword.model import Model as _OWWModel
                # a custom-trained bare-"jarvis" model beats the stock
                # "hey jarvis" one — drop it in data/ and it takes over
                custom = Path(os.path.dirname(__file__)).parent / "data" / "jarvis_custom.onnx"
                if custom.exists():
                    models = [str(custom)]
                    print(f"[wake] using custom model {custom.name}")
                else:
                    models = ["hey_jarvis_v0.1"]
                    try:
                        from openwakeword.utils import download_models as _dl
                        _dl(model_names=["hey_jarvis_v0.1"])
                    except Exception:
                        pass  # already downloaded / offline — Model() will complain if truly missing
                _oww = _OWWModel(wakeword_models=models, inference_framework="onnx")
            import io
            import wave
            with wave.open(io.BytesIO(data)) as w:
                pcm = np.frombuffer(w.readframes(w.getnframes()), dtype=np.int16)
            # ACROSS-THE-ROOM FIX: distant speech arrives quiet and the wake
            # model scores quiet audio like noise — normalize every clip to a
            # strong level so a far "Jarvis" sounds like a close-up one
            peak = int(np.abs(pcm).max()) if len(pcm) else 0
            if 0 < peak < 20000:
                pcm = (pcm.astype(np.float32) * (23000.0 / peak)).clip(-32767, 32767).astype(np.int16)
            _oww.reset()
            score = 0.0
            for i in range(0, len(pcm) - 1279, 1280):  # 80ms frames @ 16kHz
                preds = _oww.predict(pcm[i:i + 1280])
                if preds:
                    score = max(score, max(preds.values()))
        result = {"ok": True, "wake": bool(score >= 0.5), "score": round(float(score), 3)}
    except Exception as e:
        result = {"ok": False, "wake": True, "error": str(e)[:200]}

    # voiceprint gate: is this HIS voice? (auto-enrolls from confirmed wakes;
    # match=None until the profile is complete or if resemblyzer is absent)
    try:
        from jarvis.tools.voiceprint import check_and_learn
        vp = check_and_learn(data, wake_confirmed=bool(result.get("wake") and result.get("ok")))
        result["speaker"] = vp["match"]
        result["speaker_sim"] = vp["sim"]
        result["speaker_n"] = vp["n"]
    except Exception:
        pass
    return result


@app.post("/api/transcribe")
async def transcribe(audio: UploadFile = File(...)):
    """Jarvis's own ears — local Whisper first (free, vocabulary-biased),
    ElevenLabs Scribe as the cloud fallback."""
    data = await audio.read()
    if len(data) < 1500:
        return {"text": ""}  # too short to contain speech

    # local ears: faster-whisper on this laptop, primed with the room's
    # vocabulary (names, apps, teams, current movie library)
    try:
        from jarvis.tools.ears import transcribe_local
        text = await asyncio.get_event_loop().run_in_executor(None, transcribe_local, data)
        if text is not None:
            return {"text": text, "engine": "whisper"}
    except Exception:
        pass  # fall through to the cloud

    if not ELEVENLABS_API_KEY:
        return {"text": "", "error": "no ELEVENLABS_API_KEY"}
    async with httpx.AsyncClient(timeout=60) as cx:
        r = await cx.post(
            "https://api.elevenlabs.io/v1/speech-to-text",
            headers={"xi-api-key": ELEVENLABS_API_KEY},
            files={"file": (audio.filename or "audio.m4a", data, audio.content_type or "audio/mp4")},
            data={"model_id": "scribe_v1"},
        )
    if r.status_code != 200:
        return {"text": "", "error": f"stt {r.status_code}: {r.text[:200]}"}
    return {"text": r.json().get("text", "")}


class VisionRequest(BaseModel):
    image: str  # base64 encoded image
    question: str = "What is this? Describe it briefly in 2-3 sentences."


@app.post("/api/vision")
async def vision(req: VisionRequest):
    """Analyze an image with Claude's vision and return spoken response."""
    import anthropic
    from jarvis.config import ANTHROPIC_API_KEY, SYSTEM_PROMPT

    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

        # Strip data URL prefix if present
        img_data = req.image
        if "," in img_data:
            img_data = img_data.split(",", 1)[1]

        # Use Haiku for vision — much faster than Sonnet
        resp = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=150,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": "image/jpeg",
                            "data": img_data,
                        },
                    },
                    {
                        "type": "text",
                        "text": req.question,
                    },
                ],
            }],
        )
        answer = resp.content[0].text

        # Generate TTS for the answer
        audio = await generate_tts(answer)
        import urllib.parse
        encoded_text = urllib.parse.quote(answer.replace("\n", " ")[:500])

        return Response(
            content=audio,
            media_type="audio/mpeg",
            headers={"X-Jarvis-Text": encoded_text},
        )
    except Exception as e:
        return {"error": str(e)}


@app.post("/api/chat-and-speak")
async def chat_and_speak(req: ChatRequest):
    """Combined endpoint — get response and stream TTS audio."""
    response = brain.think(req.message, context=_with_tv_context(req.context))
    text = fix_pronunciation(response)

    # Try streaming from ElevenLabs — uses active voice
    if TTS_ENGINE in ("elevenlabs", "fish") and ELEVENLABS_API_KEY:
        voice = get_active_voice()
        voice_id = voice.get("voice_id", ELEVENLABS_VOICE_ID)

        async def stream_audio():
            async with httpx.AsyncClient() as client:
                async with client.stream(
                    "POST",
                    f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}/stream",
                    headers={
                        "xi-api-key": ELEVENLABS_API_KEY,
                        "Content-Type": "application/json",
                    },
                    json={
                        "text": text,
                        "model_id": "eleven_flash_v2_5",
                        "voice_settings": {
                            "stability": 0.75,
                            "similarity_boost": 0.80,
                        },
                        "optimize_streaming_latency": 4,
                    },
                    timeout=30,
                ) as resp:
                    async for chunk in resp.aiter_bytes(1024):
                        yield chunk

        import urllib.parse
        encoded_text = urllib.parse.quote(response.replace("\n", " ")[:500])
        return StreamingResponse(
            stream_audio(),
            media_type="audio/mpeg",
            headers={"X-Jarvis-Text": encoded_text},
        )

    # Fallback — non-streaming
    audio = await generate_tts(response)
    return Response(
        content=audio,
        media_type="audio/mpeg",
        headers={"X-Jarvis-Text": response.replace("\n", " ")[:500]},
    )


@app.post("/api/chat-stream")
async def chat_stream(req: ChatRequest):
    """SSE endpoint — streams text first, then audio URL."""
    import json as _json
    import urllib.parse

    async def event_stream():
        import base64

        # Step 1: Get Claude response
        response = brain.think(req.message)

        # Step 2: Send text IMMEDIATELY so frontend shows it
        yield f"data: {_json.dumps({'type': 'text', 'content': response})}\n\n"

        # Step 3: Check TTS cache first — if cached, send instantly (0ms)
        text = _truncate_for_speech(fix_pronunciation(response))
        key = _cache_key(text)
        if key in _tts_cache:
            audio_b64 = base64.b64encode(_tts_cache[key]).decode("ascii")
            yield f"data: {_json.dumps({'type': 'audio', 'content': audio_b64})}\n\n"
            yield "data: {\"type\": \"done\"}\n\n"
            return

        # Step 4: Not cached — generate full audio then send
        audio = await generate_tts(text)
        if audio:
            audio_b64 = base64.b64encode(audio).decode("ascii")
            yield f"data: {_json.dumps({'type': 'audio', 'content': audio_b64})}\n\n"

        yield "data: {\"type\": \"done\"}\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.get("/api/briefing")
async def briefing():
    """Good Morning Deagz — daily briefing. Returns text + audio."""
    import json as _json
    import base64

    # Gather all context
    dashboard_data = await dashboard()

    # Build briefing prompt
    weather = dashboard_data.get("weather", {}).get("summary", "Weather unavailable")
    forecast = dashboard_data.get("weather", {}).get("forecast", "")
    portfolio = dashboard_data.get("portfolio", {})
    bot = dashboard_data.get("bot", {})

    briefing_text = brain.think(
        f"Give me a good morning briefing. Here's the data:\n"
        f"Weather: {weather} {forecast}\n"
        f"Portfolio: balance ${portfolio.get('balance', 0):.2f}, "
        f"portfolio value ${portfolio.get('portfolio_value', 0):.2f}, "
        f"P&L ${portfolio.get('pnl', 0):.2f}, "
        f"{portfolio.get('positions', 0)} positions\n"
        f"Bot: {'running' if bot.get('running') else 'stopped'}, "
        f"{bot.get('scans', 0)} scans, {bot.get('trades', 0)} trades today\n"
        f"Keep it natural and concise — 3-4 sentences max."
    )

    text = fix_pronunciation(briefing_text)
    audio = await generate_tts(text)
    audio_b64 = base64.b64encode(audio).decode("ascii")

    return {
        "text": briefing_text,
        "audio": audio_b64,
    }


# ─── Homework Document Solver ───

from fastapi import UploadFile, File

@app.post("/api/homework/solve")
async def homework_solve(file: UploadFile = File(...)):
    """Upload a Word doc or PDF, Jarvis extracts the problems and solves them all."""
    import anthropic
    from jarvis.config import ANTHROPIC_API_KEY, get_system_prompt

    content = await file.read()
    filename = file.filename or "document"
    text = ""

    # Extract text based on file type
    if filename.endswith(".docx"):
        try:
            from docx import Document as DocxDocument
            import io
            doc = DocxDocument(io.BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            return {"error": f"Could not read Word doc: {e}"}

    elif filename.endswith(".pdf"):
        try:
            import subprocess
            import tempfile
            tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
            tmp.write(content)
            tmp.close()
            # Try pdftotext or fallback
            try:
                result = subprocess.run(["pdftotext", tmp.name, "-"], capture_output=True, text=True, timeout=10)
                text = result.stdout
            except Exception:
                text = f"[PDF file: {filename} - {len(content)} bytes. Could not extract text.]"
            Path(tmp.name).unlink(missing_ok=True)
        except Exception as e:
            return {"error": f"Could not read PDF: {e}"}

    elif filename.endswith(".txt"):
        text = content.decode("utf-8", errors="replace")

    else:
        return {"error": f"Unsupported file type: {filename}. Send .docx, .pdf, or .txt"}

    if not text.strip():
        return {"error": "Could not extract any text from the document."}

    # Send to Claude to solve
    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        resp = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[{
                "role": "user",
                "content": f"""You are a tutor. This document contains homework problems.
Solve EVERY problem step by step. For each:
1. State the problem
2. Show your work briefly
3. Give the final answer clearly

Be thorough but concise. Format with clear numbering.

DOCUMENT CONTENTS:
{text[:8000]}"""
            }],
        )
        solutions = resp.content[0].text

        # Save solutions as a Word doc
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt
            doc = DocxDocument()
            doc.add_heading(f"Homework Solutions - {filename}", level=0)
            doc.add_paragraph(datetime.datetime.now().strftime("%B %d, %Y"))
            doc.add_paragraph("")

            for line in solutions.split("\n"):
                line = line.strip()
                if not line:
                    doc.add_paragraph("")
                elif line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                else:
                    doc.add_paragraph(line)

            import datetime
            solutions_path = Path.home() / "Desktop" / f"Solutions_{filename.rsplit('.', 1)[0]}.docx"
            doc.save(str(solutions_path))
            saved_path = str(solutions_path)
        except Exception:
            saved_path = None

        # Generate TTS summary
        summary = solutions[:200] if len(solutions) > 200 else solutions
        audio = await generate_tts(f"I've solved all the problems from {filename}. The solutions document is on your desktop.")

        import base64
        return {
            "solutions": solutions,
            "saved_to": saved_path,
            "audio": base64.b64encode(audio).decode("ascii") if audio else None,
            "problem_count": solutions.count("Problem") + solutions.count("Question") + solutions.count("1."),
        }
    except Exception as e:
        return {"error": f"Failed to solve: {e}"}


# ─── Stem Player Endpoints ───

class StemRequest(BaseModel):
    query: str

@app.post("/api/stems/separate")
async def stems_separate(req: StemRequest):
    from jarvis.tools.stems import separate_song
    result = separate_song(query=req.query)
    from jarvis.tools.stems import _separation_status
    return {"message": result, "song_id": _separation_status.get("song_id", "")}

@app.get("/api/stems/{song_id}/status")
async def stems_status(song_id: str):
    import json as _json
    from jarvis.tools.stems import _separation_status, STEM_CACHE
    stem_dir = STEM_CACHE / song_id
    stems_ready = stem_dir.exists() and all(
        (stem_dir / f"{s}.mp3").exists() or (stem_dir / f"{s}.wav").exists()
        for s in ["vocals", "drums", "bass", "other"]
    )

    # Read progress file from Demucs
    progress_file = stem_dir / "progress.json"
    percent = 0
    detail = ""
    if progress_file.exists():
        try:
            prog = _json.loads(progress_file.read_text())
            percent = prog.get("percent", 0)
            detail = prog.get("detail", "")
        except Exception:
            pass

    return {
        "ready": stems_ready,
        "active": _separation_status.get("active", False),
        "progress": _separation_status.get("progress", ""),
        "song": _separation_status.get("song", ""),
        "percent": percent,
        "detail": detail,
    }

@app.get("/api/stems/library")
async def stems_library():
    """List all separated songs."""
    import json as _j
    from jarvis.tools.stems import STEM_CACHE
    songs = []
    for d in STEM_CACHE.iterdir():
        if d.is_dir() and d.name != "demucs_raw":
            has_stems = any((d / f"{s}.mp3").exists() or (d / f"{s}.wav").exists() for s in ["vocals"])
            if has_stems:
                name = d.name
                info_file = d / "info.json"
                if info_file.exists():
                    try:
                        info = _j.loads(info_file.read_text())
                        name = info.get("name", d.name)
                    except Exception:
                        pass
                songs.append({"id": d.name, "name": name})
    return {"songs": songs}

@app.get("/api/stems/{song_id}/{stem}")
async def stems_get(song_id: str, stem: str):
    from jarvis.tools.stems import STEM_CACHE
    if stem not in ("vocals", "drums", "bass", "other"):
        return Response(status_code=400)
    # Prefer MP3 (smaller, faster to load in browser)
    mp3_path = STEM_CACHE / song_id / f"{stem}.mp3"
    wav_path = STEM_CACHE / song_id / f"{stem}.wav"
    if mp3_path.exists():
        return FileResponse(str(mp3_path), media_type="audio/mpeg")
    if wav_path.exists():
        return FileResponse(str(wav_path), media_type="audio/wav")
    return Response(status_code=404)


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 3002))
    uvicorn.run(app, host="0.0.0.0", port=port)
